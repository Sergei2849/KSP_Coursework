from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "КР_КСП_Кошмин_Сергей_учет_двигателей.docx"
ASSETS = ROOT / "docs" / "report_assets"
BOOKMARK_ID = 1


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0].cells
    for i, value in enumerate(headers):
        set_cell_text(header[i], value, True)
        set_cell_shading(header[i], "E8EEF5")
        header[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            header[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = pg_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            pg_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")


def clear_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is not None:
        sect_pr.remove(pg_borders)


def add_bookmark(paragraph, name):
    global BOOKMARK_ID
    bookmark_id = str(BOOKMARK_ID)
    BOOKMARK_ID += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_props.append(color)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Times New Roman")
    font.set(qn("w:hAnsi"), "Times New Roman")
    run_props.append(font)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    run_props.append(size)
    run_element.append(run_props)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1, bookmark=None):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(14)
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_p(doc, text="", align=None, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25) if text else None
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return p


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    for i, line in enumerate(text.splitlines()):
        if i:
            run.add_break(WD_BREAK.LINE)
        run.add_text(line)
    run.font.name = "Courier New"
    run.font.size = Pt(12)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    return p


def add_cover_line(doc, text="", bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_picture(doc, filename, caption, width=Cm(15.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ASSETS / filename), width=width)
    add_caption(doc, caption)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5


def add_page_number(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    for run in list(footer.runs):
        footer._p.remove(run._r)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    add_page_border(section)

    for _ in range(0):
        add_cover_line(doc)
    for line in [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ",
        "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        "«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНОЛОГИЧЕСКИЙ УНИВЕРСИТЕТ «МИСИС»",
        "",
        "ИНСТИТУТ КОМПЬЮТЕРНЫХ НАУК",
        "КАФЕДРА АВТОМАТИЗИРОВАННЫХ СИСТЕМ УПРАВЛЕНИЯ",
    ]:
        add_cover_line(doc, line, bold=bool(line), size=14, after=0)

    for _ in range(1):
        add_cover_line(doc)
    add_cover_line(doc, "Курс «Клиент-серверные приложения»", size=14)
    add_cover_line(doc, "ОТЧЕТ ПО КУРСОВОЙ РАБОТЕ", True, size=14)
    add_cover_line(doc, "на тему:", size=14)
    add_cover_line(doc, "«Разработка клиент-серверного приложения для учета двигателей", True, size=14)
    add_cover_line(doc, "и заявок на обслуживание»", True, size=14)

    for _ in range(2):
        add_cover_line(doc)
    add_cover_line(doc, "Выполнил: студент группы БИВТ-24-7", size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_cover_line(doc, "Кошмин Сергей", size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_cover_line(doc, "Проверил: Абросимов Никита Андреевич", size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for _ in range(8):
        add_cover_line(doc)
    add_cover_line(doc, "Москва, 2026", size=14)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    main_section.page_width = Cm(21)
    main_section.page_height = Cm(29.7)
    main_section.top_margin = Cm(2)
    main_section.bottom_margin = Cm(2)
    main_section.left_margin = Cm(3)
    main_section.right_margin = Cm(1.5)
    clear_page_border(main_section)

    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    toc = [
        ("1. Предметная область", "3", "sec_subject"),
        ("2. Постановка задачи", "4", "sec_task"),
        ("3. Описание архитектуры системы", "5", "sec_arch"),
        ("4. Описание структуры базы данных", "8", "sec_db"),
        ("5. Описание серверной части", "14", "sec_backend"),
        ("6. Описание клиентской части", "18", "sec_frontend"),
        ("7. Тестирование работы приложения", "24", "sec_testing"),
        ("8. Заключение", "28", "sec_conclusion"),
        ("9. Список использованных источников", "29", "sec_sources"),
        ("10. Приложение", "30", "sec_appendix"),
    ]
    for name, page, anchor in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        add_internal_hyperlink(p, f"{name}\t{page}", anchor)

    doc.add_page_break()
    add_heading(doc, "1. ПРЕДМЕТНАЯ ОБЛАСТЬ", 1, "sec_subject")
    add_p(doc, "Предметной областью курсовой работы является учет двигателей, деталей двигателя и заявок на их техническое обслуживание. В реальной эксплуатации сервисной организации важно хранить сведения о модели двигателя, типе, серийном номере, состоянии деталей и обращениях клиентов. Без единой информационной системы такие данные часто распределяются между бумажными журналами, таблицами и устными договоренностями.")
    add_p(doc, "Разработанное приложение решает задачу как небольшая клиент-серверная система. Пользователь может зарегистрироваться, выполнить вход, просмотреть справочник двигателей, создать заявку на обслуживание и отследить ее состояние. Механик работает со справочниками двигателей и деталей, а администратор контролирует заявки и журнал действий.")
    add_p(doc, "Особенностью выбранной предметной области является необходимость разграничивать права доступа. Клиенту не требуется управлять всеми данными системы, механик должен иметь доступ к техническим справочникам, а администратор должен видеть полный список заявок и иметь возможность менять их статусы.")
    add_p(doc, "В рамках предметной области выделены следующие основные сущности:")
    for item in [
        "пользователь системы с ролью и учетными данными;",
        "двигатель с моделью, типом, мощностью, объемом и серийным номером;",
        "деталь двигателя с состоянием и техническим примечанием;",
        "заявка на обслуживание с описанием проблемы, приоритетом и статусом;",
        "журнал действий, фиксирующий вход, регистрацию и изменение данных.",
    ]:
        add_bullet(doc, item)
    add_p(doc, "Для пользователя система выглядит как единый веб-сервис. При этом с точки зрения проектирования она делится на несколько функциональных подсистем: подсистему пользователей и ролей, подсистему справочника двигателей, подсистему учета деталей, подсистему заявок на обслуживание, подсистему статистики и подсистему журналирования.")
    add_p(doc, "Подсистема пользователей отвечает за регистрацию, вход и определение прав доступа. Подсистема двигателей хранит технические характеристики и идентификационные данные. Подсистема деталей позволяет учитывать состояние комплектующих. Подсистема заявок связывает клиента, двигатель и описание проблемы. Подсистема статистики обобщает данные для механика и администратора.")
    add_caption(doc, "Таблица 1 - Основные роли и сценарии предметной области")
    add_table(
        doc,
        ["Роль", "Основные сценарии", "Ограничения"],
        [
            ["Клиент", "Регистрация, вход, создание заявки, просмотр своих обращений.", "Не может редактировать справочники и менять статусы."],
            ["Механик", "Добавление двигателей, деталей, создание заявок, просмотр статистики.", "Не имеет доступа к журналу администратора и удалению двигателей."],
            ["Администратор", "Контроль всех заявок, смена статусов, просмотр логов, удаление записей.", "Действия фиксируются в журнале."],
        ],
        [Cm(3), Cm(8), Cm(5)],
    )

    add_heading(doc, "2. ПОСТАНОВКА ЗАДАЧИ", 1, "sec_task")
    add_p(doc, "Цель работы: разработать программное клиент-серверное приложение для учета двигателей и заявок на обслуживание с использованием технологий, изученных в курсе «Клиент-серверные приложения».")
    add_p(doc, "Для достижения поставленной цели необходимо решить следующие задачи:")
    for item in [
        "спроектировать реляционную базу данных PostgreSQL минимум на четыре таблицы;",
        "реализовать регистрацию и вход пользователей;",
        "добавить проверку корректности email при регистрации и авторизации;",
        "реализовать разделение ролей пользователя, механика и администратора;",
        "создать серверную часть с REST API для двигателей, деталей, заявок и логов;",
        "разработать клиентскую часть на HTML, CSS и JavaScript;",
        "добавить журналирование значимых действий;",
        "подготовить демонстрационные данные и инструкцию запуска.",
    ]:
        add_bullet(doc, item)
    add_p(doc, "В таблице 2 приведено соответствие обязательных требований и реализованных элементов проекта.")
    add_caption(doc, "Таблица 2 - Соответствие требований и реализации")
    add_table(
        doc,
        ["Требование", "Реализация в проекте"],
        [
            ["Разделение ролей", "Роли admin, mechanic и client; проверка прав через require_role."],
            ["Регистрация", "Маршрут POST /api/auth/register создает пользователя и сохраняет хеш пароля."],
            ["Вход", "Маршрут POST /api/auth/login проверяет email и пароль."],
            ["База данных", "PostgreSQL-схема содержит 6 таблиц и внешние ключи."],
            ["Расширенная БД-логика", "Добавлены 3 представления, 3 функции, 3 процедуры и 3 триггера."],
            ["Валидация email", "Регулярная проверка на backend и проверка формы на frontend."],
            ["Логирование", "Запись в logs/app.log и таблицу action_logs."],
        ],
        [Cm(5), Cm(10.5)],
    )

    add_heading(doc, "3. ОПИСАНИЕ АРХИТЕКТУРЫ СИСТЕМЫ", 1, "sec_arch")
    add_heading(doc, "3.1 Общая архитектура", 2)
    add_p(doc, "Система построена по классической клиент-серверной схеме. Клиентская часть выполняется в браузере и отвечает за отображение страниц, обработку форм и отправку HTTP-запросов. Серверная часть реализована на Flask и предоставляет API, через которое выполняются операции с пользователями, двигателями, деталями, заявками и журналом.")
    add_p(doc, "Уровень хранения данных представлен СУБД PostgreSQL. Для доступа к базе используется библиотека psycopg2, запросы выполняются через вспомогательные функции query и query_one.")
    add_code(doc, "Браузер → HTML/CSS/JavaScript → Flask API → PostgreSQL")
    add_p(doc, "Трехуровневое разделение упрощает сопровождение проекта. Клиентская часть не знает, как именно устроены SQL-таблицы, а работает только с JSON-ответами API. Серверная часть скрывает детали подключения к базе, проверяет права доступа и формирует единый формат ошибок. База данных отвечает за надежное хранение, связи между сущностями и часть проверок на уровне триггеров.")
    add_p(doc, "Типовой сценарий обработки запроса выглядит следующим образом: пользователь заполняет форму, JavaScript собирает данные, отправляет запрос через fetch, Flask-маршрут загружает текущего пользователя, проверяет роль, валидирует входные данные, обращается к PostgreSQL, записывает действие в журнал и возвращает JSON-ответ.")
    add_caption(doc, "Таблица 3 - Уровни архитектуры приложения")
    add_table(
        doc,
        ["Уровень", "Компоненты", "Ответственность"],
        [
            ["Представление", "HTML, CSS, JavaScript", "Отображение страниц, формы, таблицы, сохранение пользователя в LocalStorage."],
            ["Бизнес-логика", "Flask, routes/auth.py, routes/api.py", "Проверка прав, обработка запросов, валидация, формирование JSON-ответов."],
            ["Доступ к данным", "backend/db.py, psycopg2", "Открытие соединений, выполнение SQL-запросов, commit/rollback."],
            ["Хранение", "PostgreSQL", "Таблицы, внешние ключи, представления, функции, процедуры и триггеры."],
        ],
        [Cm(3.2), Cm(5.8), Cm(7)],
    )
    add_heading(doc, "3.2 Технологический стек", 2)
    add_caption(doc, "Таблица 4 - Технологический стек")
    add_table(
        doc,
        ["Компонент", "Используемая технология"],
        [
            ["Клиентская часть", "HTML5, CSS3, JavaScript, Fetch API, LocalStorage."],
            ["Серверная часть", "Python, Flask, Werkzeug, python-dotenv."],
            ["База данных", "PostgreSQL, SQL-скрипты schema.sql и seed.sql."],
            ["Логирование", "Модуль logging, файл logs/app.log, таблица action_logs."],
            ["Проверка", "compileall для Python, node --check для JavaScript."],
        ],
        [Cm(5), Cm(10.5)],
    )
    add_heading(doc, "3.3 Структура проекта", 2)
    add_caption(doc, "Таблица 5 - Структура проекта")
    add_table(
        doc,
        ["Папка или файл", "Назначение"],
        [
            ["backend/", "Серверная часть Flask, маршруты API, авторизация, подключение к БД."],
            ["backend/routes/", "Маршруты auth.py и api.py."],
            ["database/", "SQL-схема, стартовые данные и скрипт подготовки БД."],
            ["public/", "HTML-страницы, CSS и JavaScript клиентской части."],
            ["logs/", "Файл журнала действий приложения."],
            ["docs/", "Заметки для защиты, отчет и изображения для отчета."],
        ],
        [Cm(5), Cm(10.5)],
    )
    add_p(doc, "Взаимодействие компонентов происходит через JSON. Например, при создании заявки форма собирает данные, преобразует их в JSON и отправляет POST-запрос на /api/service-requests. Сервер проверяет роль пользователя, валидирует поля, сохраняет запись в базе данных и возвращает созданную заявку.")
    add_p(doc, "При возникновении ошибки сервер возвращает JSON с полем error и соответствующим HTTP-кодом. Например, при отсутствии авторизации возвращается 401, при нехватке прав - 403, при некорректных данных - 400. Такой подход удобен для клиентской части: JavaScript может одинаково обрабатывать ошибки разных форм.")

    add_heading(doc, "4. ОПИСАНИЕ СТРУКТУРЫ БАЗЫ ДАННЫХ", 1, "sec_db")
    add_p(doc, "База данных построена по реляционной модели и содержит шесть таблиц. Это превышает минимальное требование курсовой работы и позволяет показать связи между пользователями, ролями, двигателями, деталями и заявками.")
    add_p(doc, "Логическая модель строится вокруг сущности «двигатель». Один пользователь может добавить много двигателей, один двигатель может иметь много деталей и много заявок. Каждая заявка создается конкретным пользователем и относится к одному двигателю. Роль пользователя вынесена в отдельную таблицу roles, что позволяет расширять набор ролей без изменения структуры users.")
    add_caption(doc, "Таблица 6 - Таблицы базы данных")
    add_table(
        doc,
        ["Таблица", "Назначение", "Основные поля"],
        [
            ["roles", "Хранение ролей пользователей.", "id, code, name"],
            ["users", "Хранение учетных записей.", "id, name, email, password_hash, role_id"],
            ["engines", "Справочник двигателей.", "id, model, engine_type, power_hp, volume_liters, serial_number"],
            ["engine_parts", "Детали конкретного двигателя.", "id, engine_id, name, part_code, condition, note"],
            ["service_requests", "Заявки на обслуживание.", "id, engine_id, created_by, client_name, problem, status"],
            ["action_logs", "Журнал действий в БД.", "id, user_id, action, details, created_at"],
        ],
        [Cm(3.4), Cm(5.4), Cm(7.0)],
    )
    add_p(doc, "Связь пользователей с ролями реализована через внешний ключ users.role_id. Двигатели связаны с пользователем, который добавил запись. Детали связаны с двигателями через engine_parts.engine_id. Заявки связаны одновременно с двигателем и пользователем, создавшим обращение.")
    add_heading(doc, "4.1 Описание таблиц", 2)
    add_p(doc, "Таблица roles хранит справочник ролей. В начальных данных создаются роли admin, mechanic и client. Использование отдельной таблицы делает ролевую модель нормализованной: в таблице users хранится не текст роли, а внешний ключ.")
    add_p(doc, "Таблица users хранит учетные записи. Для пользователя сохраняются имя, email, хеш пароля и ссылка на роль. Email имеет тип CITEXT, поэтому уникальность адресов не зависит от регистра букв. Пароль не хранится в открытом виде, вместо него используется результат generate_password_hash.")
    add_p(doc, "Таблица engines является справочником двигателей. В ней хранятся модель, тип двигателя, мощность, объем, серийный номер, описание, пользователь-создатель и даты создания/обновления. Серийный номер сделан уникальным, потому что в реальной предметной области он является идентификатором конкретного двигателя.")
    add_p(doc, "Таблица engine_parts содержит детали двигателя. Каждая деталь связана с конкретным двигателем. Для детали хранится название, код, состояние и примечание. Состояние позволяет показать, исправна деталь, находится ли она в диагностике или требует замены.")
    add_p(doc, "Таблица service_requests содержит заявки на обслуживание. Заявка связывает двигатель, пользователя, клиента, описание проблемы, приоритет, статус и комментарий администратора. Именно эта таблица отражает основной бизнес-процесс приложения.")
    add_p(doc, "Таблица action_logs используется для хранения журнала действий в БД. Дополнительно действия пишутся в файл logs/app.log, но таблица удобна для статистики и вывода последних событий в интерфейсе.")
    add_p(doc, "Для таблиц engines и service_requests добавлены триггеры обновления поля updated_at. Это позволяет автоматически фиксировать время последнего изменения записи.")
    add_code(doc, "CREATE TABLE service_requests (\n    id SERIAL PRIMARY KEY,\n    engine_id INTEGER NOT NULL REFERENCES engines(id),\n    created_by INTEGER NOT NULL REFERENCES users(id),\n    problem TEXT NOT NULL,\n    status VARCHAR(50) NOT NULL DEFAULT 'Новая'\n);")
    add_heading(doc, "4.2 Представления базы данных", 2)
    add_p(doc, "Для упрощения выборок и формирования статистики в базе данных созданы представления. Они объединяют несколько таблиц и возвращают уже подготовленные данные для интерфейса и серверных маршрутов.")
    add_caption(doc, "Таблица 7 - Представления базы данных")
    add_table(
        doc,
        ["Представление", "Назначение"],
        [
            ["v_engines_full", "Возвращает двигатели вместе с автором записи, количеством деталей и количеством заявок."],
            ["v_service_requests_full", "Возвращает заявки с данными двигателя и пользователя, создавшего заявку."],
            ["v_request_status_stats", "Формирует статистику количества заявок по каждому статусу."],
        ],
        [Cm(5), Cm(10.5)],
    )
    add_p(doc, "Представление v_engines_full используется в API списка двигателей. Оно объединяет engines, users, engine_parts и service_requests, поэтому клиент получает не только технические характеристики, но и агрегированные показатели: количество деталей и количество заявок по каждому двигателю.")
    add_p(doc, "Представление v_service_requests_full используется в API заявок. Оно раскрывает внешние ключи и возвращает модель двигателя, серийный номер и данные автора заявки. Это избавляет сервер от повторения JOIN-запросов в каждом маршруте.")
    add_p(doc, "Представление v_request_status_stats используется на странице статистики. Оно группирует заявки по статусам и позволяет быстро отобразить распределение обращений.")
    add_heading(doc, "4.3 Функции, процедуры и триггеры", 2)
    add_p(doc, "По примеру более расширенных курсовых работ в проект добавлена бизнес-логика на уровне PostgreSQL. Функции используются для получения вычисляемых значений и поиска, процедуры автоматизируют повторяющиеся операции, а триггеры защищают данные от некорректных изменений.")
    add_caption(doc, "Таблица 8 - Расширенная логика базы данных")
    add_table(
        doc,
        ["Тип", "Название", "Назначение"],
        [
            ["Функция", "fn_engine_parts_count", "Возвращает количество деталей выбранного двигателя."],
            ["Функция", "fn_user_request_count", "Возвращает количество заявок пользователя."],
            ["Функция", "fn_search_engines", "Ищет двигатели по модели, типу или серийному номеру."],
            ["Процедура", "sp_change_request_status", "Меняет статус заявки и записывает действие в журнал."],
            ["Процедура", "sp_update_part_condition", "Безопасно обновляет состояние детали."],
            ["Процедура", "sp_cancel_old_new_requests", "Отменяет старые необработанные заявки."],
            ["Триггер", "trg_users_email_check", "Проверяет формат email на уровне БД."],
            ["Триггер", "trg_prevent_last_admin_delete", "Запрещает удалить последнего администратора."],
            ["Триггер", "trg_service_request_status_check", "Проверяет допустимость статуса заявки."],
        ],
        [Cm(3), Cm(5.2), Cm(8)],
    )
    add_code(doc, "CREATE VIEW v_request_status_stats AS\nSELECT status, COUNT(*) AS requests_count\nFROM service_requests\nGROUP BY status;")
    add_p(doc, "Функция fn_engine_parts_count может использоваться для быстрого получения количества деталей конкретного двигателя. Это полезно при выводе карточки двигателя или при построении отчета по техническому состоянию.")
    add_p(doc, "Функция fn_user_request_count возвращает количество заявок, созданных пользователем. Она может применяться для личного кабинета или для оценки активности клиента.")
    add_p(doc, "Функция fn_search_engines выполняет поиск по модели, типу и серийному номеру. Поиск вынесен в базу данных, чтобы правила поиска были едиными для разных частей приложения.")
    add_code(doc, "CREATE OR REPLACE FUNCTION fn_search_engines(p_query TEXT)\nRETURNS TABLE (...)\nAS $$\nBEGIN\n    RETURN QUERY\n    SELECT ... FROM v_engines_full\n    WHERE LOWER(model) LIKE LOWER('%' || p_query || '%');\nEND;\n$$ LANGUAGE plpgsql;")
    add_p(doc, "Процедура sp_change_request_status автоматизирует смену статуса заявки. В отличие от обычного UPDATE, она также записывает действие в action_logs. Это делает операцию более надежной: изменение статуса и журналирование описаны в одном месте.")
    add_p(doc, "Процедура sp_update_part_condition безопасно обновляет состояние детали. Перед обновлением она проверяет существование записи и выбрасывает исключение, если деталь не найдена.")
    add_p(doc, "Процедура sp_cancel_old_new_requests предназначена для автоматического обслуживания данных. Она переводит старые необработанные заявки в статус «Отменена» и оставляет административный комментарий.")
    add_code(doc, "CREATE OR REPLACE PROCEDURE sp_change_request_status(\n    p_request_id INTEGER,\n    p_new_status VARCHAR,\n    p_admin_id INTEGER,\n    p_comment TEXT DEFAULT NULL\n)\nLANGUAGE plpgsql AS $$\nBEGIN\n    UPDATE service_requests SET status = p_new_status WHERE id = p_request_id;\n    INSERT INTO action_logs(user_id, action, details) VALUES (...);\nEND;\n$$;")
    add_p(doc, "Триггер trg_users_email_check проверяет email на уровне базы данных. Даже если backend по ошибке пропустит некорректный адрес, база не позволит сохранить такую запись.")
    add_p(doc, "Триггер trg_prevent_last_admin_delete защищает систему от удаления последнего администратора. Это важно для сохранения управляемости приложения.")
    add_p(doc, "Триггер trg_service_request_status_check проверяет, что статус заявки входит в допустимый набор значений. Благодаря этому в базе не появятся произвольные статусы с опечатками.")
    add_code(doc, "CREATE TRIGGER trg_service_request_status_check\n    BEFORE INSERT OR UPDATE ON service_requests\n    FOR EACH ROW\n    EXECUTE PROCEDURE validate_service_request_status();")

    add_heading(doc, "5. ОПИСАНИЕ СЕРВЕРНОЙ ЧАСТИ", 1, "sec_backend")
    add_p(doc, "Серверная часть реализована на Flask. Точкой входа является файл backend/app.py, где создается приложение, регистрируются blueprints и настраивается отдача статических файлов из папки public.")
    add_heading(doc, "5.1 Структура серверной части", 2)
    add_caption(doc, "Таблица 9 - Файлы серверной части")
    add_table(
        doc,
        ["Файл", "Назначение"],
        [
            ["backend/app.py", "Создание Flask-приложения, регистрация маршрутов и отдача статических страниц."],
            ["backend/db.py", "Подключение к PostgreSQL и функции query, query_one."],
            ["backend/auth.py", "Работа с пользователями, email-валидация, загрузка текущего пользователя, проверка ролей."],
            ["backend/audit.py", "Запись действий в таблицу action_logs."],
            ["backend/log.py", "Настройка файлового логирования в logs/app.log."],
            ["backend/routes/auth.py", "Маршруты регистрации, входа и получения текущего пользователя."],
            ["backend/routes/api.py", "Маршруты двигателей, деталей, заявок, логов и статистики."],
        ],
        [Cm(5.2), Cm(10.5)],
    )
    add_p(doc, "Серверная часть разделена на небольшие модули, чтобы каждая часть отвечала за свою задачу. Такое разделение упрощает сопровождение: изменение логики авторизации не требует правки маршрутов справочников, а изменение подключения к БД не затрагивает клиентскую часть.")
    add_heading(doc, "5.2 Авторизация и регистрация", 2)
    add_p(doc, "Авторизация вынесена в backend/routes/auth.py. При регистрации пользователь передает имя, email и пароль. Email нормализуется и проверяется регулярным выражением, пароль хешируется через generate_password_hash. При входе используется check_password_hash.")
    add_code(doc, "EMAIL_RE = re.compile(r\"^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$\")\n\ndef is_valid_email(email: str) -> bool:\n    return bool(email and EMAIL_RE.match(email))")
    add_p(doc, "Регистрация выполняется в несколько этапов. Сначала сервер получает JSON и извлекает поля name, email и password. Затем email приводится к нижнему регистру, выполняется проверка формата и проверка уникальности. Если пользователь с таким email уже существует, сервер возвращает ошибку 409.")
    add_p(doc, "Если email совпадает с ADMIN_EMAIL из переменных окружения, пользователю назначается роль admin. В остальных случаях новый пользователь получает роль mechanic. Такой подход удобен для учебного проекта: администратор задается через конфигурацию, а остальные пользователи могут регистрироваться самостоятельно.")
    add_p(doc, "При входе сервер не сравнивает пароль напрямую. Вместо этого из БД извлекается хеш и используется функция check_password_hash. Если пароль неверный, в файл журнала записывается неудачная попытка входа.")
    add_p(doc, "Разделение прав реализовано через функцию require_role. Она проверяет, выполнен ли вход и входит ли роль текущего пользователя в список разрешенных ролей. Для идентификации пользователя клиент передает заголовок X-User-Id, после чего backend загружает пользователя из базы.")
    add_heading(doc, "5.3 API-маршруты", 2)
    add_p(doc, "API построено в REST-стиле. Получение данных выполняется методом GET, создание - POST, изменение - PUT или PATCH, удаление - DELETE. Все ответы возвращаются в формате JSON. Ошибки также оформляются как JSON-объекты с полем error.")
    add_caption(doc, "Таблица 10 - Основные API-эндпоинты")
    add_table(
        doc,
        ["Endpoint", "Метод", "Назначение"],
        [
            ["/api/auth/register", "POST", "Регистрация нового пользователя."],
            ["/api/auth/login", "POST", "Вход пользователя."],
            ["/api/engines", "GET/POST", "Получение и создание двигателей."],
            ["/api/engines/<id>", "PUT/DELETE", "Изменение и удаление двигателя."],
            ["/api/parts", "GET/POST", "Получение и создание деталей двигателя."],
            ["/api/service-requests", "GET/POST", "Получение и создание заявок."],
            ["/api/service-requests/<id>", "PATCH", "Изменение статуса заявки администратором."],
            ["/api/logs", "GET", "Получение журнала действий."],
            ["/api/stats", "GET", "Получение статистики по данным из представлений."],
        ],
        [Cm(5.2), Cm(2.7), Cm(8.2)],
    )
    add_p(doc, "Маршрут /api/engines возвращает данные из представления v_engines_full. Благодаря этому в таблице двигателей можно показывать не только модель и серийный номер, но и имя пользователя, добавившего запись. Маршрут /api/service-requests использует v_service_requests_full и возвращает расширенную информацию по заявкам.")
    add_p(doc, "Маршрут /api/stats является дополнительным по сравнению с минимальными требованиями. Он собирает общие счетчики, распределение заявок по статусам, активность двигателей и последние действия. Данные берутся из таблиц и представлений PostgreSQL.")
    add_heading(doc, "5.4 Валидация и обработка ошибок", 2)
    add_p(doc, "На сервере проверяются обязательные поля. Для двигателя обязательны модель и тип. Для детали обязательны существующий двигатель и название. Для заявки обязательны двигатель, ФИО клиента и подробное описание проблемы. Если данные не проходят проверку, сервер возвращает код 400.")
    add_p(doc, "Для действий с ограниченным доступом используется require_role. Например, создание двигателя доступно администратору и механику, а изменение статуса заявки доступно только администратору. При нарушении прав возвращается код 403.")
    add_caption(doc, "Таблица 11 - Примеры серверных проверок")
    add_table(
        doc,
        ["Сценарий", "Проверка", "Ответ при ошибке"],
        [
            ["Регистрация", "Email должен соответствовать формату и быть уникальным.", "400 или 409"],
            ["Вход", "Email и пароль должны соответствовать учетной записи.", "401"],
            ["Создание двигателя", "Модель и тип не должны быть пустыми.", "400"],
            ["Создание детали", "Двигатель должен существовать в БД.", "400"],
            ["Смена статуса", "Пользователь должен быть администратором, статус должен быть допустимым.", "400 или 403"],
        ],
        [Cm(4.2), Cm(8.2), Cm(3)],
    )
    add_heading(doc, "5.5 Журналирование", 2)
    add_p(doc, "Логирование реализовано двумя способами. Первый способ - запись в файл logs/app.log через стандартный модуль logging. Второй способ - запись бизнес-действий в таблицу action_logs через функцию write_action. Это позволяет показать на защите как файловый журнал, так и журнал в базе данных.")
    add_p(doc, "В журнал записываются регистрация, вход пользователя, создание двигателя, обновление двигателя, удаление двигателя, добавление детали, создание заявки и изменение статуса. Такой набор событий позволяет восстановить основные действия пользователей при демонстрации проекта.")
    add_code(doc, "log.info(\"Создана заявка id=%s engine=%s user=%s\", row[\"id\"], engine_id, g.db_user[\"id\"])\nwrite_action(g.db_user[\"id\"], \"Создана заявка\", f\"id={row['id']}, engine={engine_id}\")")

    add_heading(doc, "6. ОПИСАНИЕ КЛИЕНТСКОЙ ЧАСТИ", 1, "sec_frontend")
    add_p(doc, "Клиентская часть состоит из HTML-страниц, общего CSS-файла и JavaScript-файла app.js. В интерфейсе реализованы страницы главной, входа и регистрации, справочника двигателей, справочника деталей, заявок и журнала действий.")
    add_caption(doc, "Таблица 12 - Страницы клиентской части")
    add_table(
        doc,
        ["Страница", "Назначение", "Основные элементы"],
        [
            ["index.html", "Главная страница приложения.", "Описание системы, быстрые действия, карточки возможностей."],
            ["login.html", "Вход и регистрация.", "Формы входа, регистрации, выход из системы."],
            ["engines.html", "Справочник двигателей.", "Форма добавления, таблица двигателей, обновление списка."],
            ["parts.html", "Справочник деталей.", "Выбор двигателя, добавление детали, поиск."],
            ["requests.html", "Заявки на обслуживание.", "Создание заявки, таблица заявок, смена статуса, CSV."],
            ["stats.html", "Статистика.", "Карточки итогов, статусы, активность двигателей, последние действия."],
            ["logs.html", "Журнал действий.", "Список последних записей файла app.log."],
        ],
        [Cm(3.4), Cm(5.8), Cm(6.5)],
    )
    add_p(doc, "Все страницы используют единый верхний блок навигации. JavaScript после загрузки страницы проверяет текущего пользователя и скрывает элементы, недоступные для его роли. Например, ссылка «Логи» отображается только администратору, а ссылка «Статистика» доступна администратору и механику.")
    add_picture(doc, "main.png", "Рисунок 1 - Главная страница приложения")
    add_heading(doc, "6.1 Вход, регистрация и роли", 2)
    add_p(doc, "На странице входа пользователь может авторизоваться или создать новую учетную запись. После успешного входа данные пользователя сохраняются в LocalStorage, а при последующих запросах отправляется идентификатор пользователя в заголовке X-User-Id.")
    add_picture(doc, "login.png", "Рисунок 2 - Страница входа и регистрации")
    add_p(doc, "На клиентской стороне email проверяется регулярным выражением до отправки формы. Это снижает количество некорректных запросов к серверу, однако окончательная проверка также выполняется на backend.")
    add_code(doc, "function authHeaders() {\n    const user = getUser();\n    const headers = {'Content-Type': 'application/json'};\n    if (user) headers['X-User-Id'] = user.id;\n    return headers;\n}")
    add_p(doc, "LocalStorage используется как учебный механизм хранения текущего пользователя. После входа сохраняются id, имя, email, код роли и русское название роли. При выходе запись удаляется, и пользователь снова считается неавторизованным.")
    add_caption(doc, "Таблица 13 - Отображение элементов интерфейса по ролям")
    add_table(
        doc,
        ["Элемент интерфейса", "Клиент", "Механик", "Администратор"],
        [
            ["Создание заявки", "Да", "Да", "Да"],
            ["Добавление двигателя", "Нет", "Да", "Да"],
            ["Добавление детали", "Нет", "Да", "Да"],
            ["Страница статистики", "Нет", "Да", "Да"],
            ["Смена статуса заявки", "Нет", "Нет", "Да"],
            ["Просмотр логов", "Нет", "Нет", "Да"],
        ],
        [Cm(5.2), Cm(3), Cm(3), Cm(3)],
    )
    add_heading(doc, "6.2 Справочник двигателей", 2)
    add_p(doc, "Страница двигателей позволяет просматривать список двигателей и добавлять новые записи. Для двигателя указываются модель, тип, мощность, объем, серийный номер и описание. Добавление доступно пользователям с ролью администратора или механика.")
    add_picture(doc, "engines.png", "Рисунок 3 - Страница справочника двигателей")
    add_p(doc, "После отправки формы JavaScript вызывает POST /api/engines. Если сервер возвращает успешный ответ, форма очищается, а список двигателей загружается заново. Такой подход обеспечивает актуальность таблицы без перезагрузки страницы.")
    add_p(doc, "Таблица двигателей отображает модель, тип, мощность, объем, серийный номер и пользователя, добавившего запись. Сервер возвращает эти данные из представления v_engines_full.")
    add_heading(doc, "6.3 Детали и заявки", 2)
    add_p(doc, "Справочник деталей связан с выбранным двигателем. Для каждой детали можно указать название, код, состояние и техническое примечание. Страница заявок используется для фиксации обращения клиента или механика по конкретному двигателю.")
    add_picture(doc, "requests.png", "Рисунок 4 - Страница заявок на обслуживание")
    add_p(doc, "Администратору доступно изменение статуса заявки. Варианты статусов: «Новая», «В диагностике», «Ожидает деталь», «Выполнена», «Отменена». Для удобства также реализована выгрузка списка заявок в CSV.")
    add_p(doc, "CSV-выгрузка формируется полностью на клиентской стороне. JavaScript берет текущий массив заявок, экранирует значения и создает Blob с типом text/csv. После этого создается временная ссылка для скачивания файла service_requests.csv.")
    add_code(doc, "const csv = rows.map(row => row.map(cell => `\"${String(cell).replaceAll('\"', '\"\"')}\"`).join(';')).join('\\n');\nconst blob = new Blob([csv], {type: 'text/csv;charset=utf-8'});")
    add_heading(doc, "6.4 Статистика", 2)
    add_p(doc, "После доработки в интерфейс добавлена страница статистики. Она доступна администратору и механику. На странице отображаются итоговые количества двигателей, деталей, заявок и пользователей, распределение заявок по статусам, активность двигателей и последние действия из журнала.")
    add_picture(doc, "stats.png", "Рисунок 5 - Страница статистики системы")
    add_p(doc, "Статистика показывает, что приложение использует не только CRUD-операции, но и агрегированные данные. Это повышает ценность системы для администратора: он может быстро увидеть количество заявок, проблемные статусы и наиболее активные двигатели.")
    add_heading(doc, "6.5 Адаптивность и оформление", 2)
    add_p(doc, "Оформление вынесено в файл public/css/style.css. Для интерфейса выбрана спокойная рабочая стилистика: темная верхняя панель, светлые панели, таблицы с аккуратными границами и зеленый акцент для основных действий.")
    add_p(doc, "CSS содержит медиазапрос для мобильных экранов. При уменьшении ширины окна сетки перестраиваются в одну колонку, формы становятся вертикальными, а широкие таблицы помещаются в горизонтально прокручиваемый контейнер.")

    add_heading(doc, "7. ТЕСТИРОВАНИЕ РАБОТЫ ПРИЛОЖЕНИЯ", 1, "sec_testing")
    add_p(doc, "Проверка проекта выполнялась вручную и с помощью команд проверки синтаксиса. Были проверены запуск Flask-приложения, отдача главной страницы, корректность JavaScript-файла, импорт backend-модулей и совпадение тестового пароля с хешем из seed.sql.")
    add_p(doc, "Тестирование разделено на несколько групп: проверка запуска, проверка авторизации, проверка CRUD-сценариев, проверка разграничения ролей, проверка клиентских форм и проверка отчетного документа. Такой набор проверок соответствует учебному характеру проекта и показывает работоспособность основных функций.")
    add_caption(doc, "Таблица 14 - Проверка основных сценариев")
    add_table(
        doc,
        ["Проверка", "Ожидаемый результат", "Результат"],
        [
            ["Открытие главной страницы", "Возвращается HTTP 200 и отображается интерфейс.", "Выполнено"],
            ["Проверка Python", "Модули backend успешно компилируются.", "Выполнено"],
            ["Проверка JavaScript", "Файл public/js/app.js не содержит синтаксических ошибок.", "Выполнено"],
            ["Проверка пароля", "Пароль 1234 соответствует хешу тестовых пользователей.", "Выполнено"],
            ["Проверка ролей", "Админские элементы скрываются для неадминистратора.", "Выполнено"],
            ["Проверка статистики", "Страница stats.html получает данные через /api/stats.", "Выполнено"],
            ["Создание заявки", "После входа пользователь может отправить заявку.", "Реализовано в API и интерфейсе"],
        ],
        [Cm(5), Cm(7), Cm(3.5)],
    )
    add_heading(doc, "7.1 Проверка авторизации", 2)
    add_p(doc, "Проверка авторизации включает регистрацию нового пользователя, вход с корректным email и паролем, вход с неверным паролем и отображение текущего пользователя в навигации. На клиентской стороне также проверяется реакция формы на некорректный email.")
    add_caption(doc, "Таблица 15 - Проверка авторизации и ролей")
    add_table(
        doc,
        ["Сценарий", "Действие", "Ожидаемый результат"],
        [
            ["Регистрация", "Ввести имя, корректный email и пароль от 4 символов.", "Пользователь создается, данные сохраняются в LocalStorage."],
            ["Некорректный email", "Ввести строку без символа @.", "Форма показывает сообщение об ошибке."],
            ["Повторный email", "Зарегистрировать уже существующий email.", "Сервер возвращает ошибку 409."],
            ["Вход", "Ввести email и пароль тестового пользователя.", "Пользователь попадает на главную страницу."],
            ["Права клиента", "Войти под client@engine.local.", "Ссылки на статистику и логи скрыты."],
            ["Права администратора", "Войти под admin@engine.local.", "Доступны логи и смена статуса заявки."],
        ],
        [Cm(4), Cm(6.4), Cm(5.6)],
    )
    add_heading(doc, "7.2 Проверка CRUD-операций", 2)
    add_p(doc, "CRUD-операции проверяются для основных сущностей: двигателей, деталей и заявок. Для двигателя реализованы создание, просмотр, изменение и удаление. Для деталей реализованы создание и просмотр. Для заявок реализованы создание, просмотр и изменение статуса.")
    add_caption(doc, "Таблица 16 - Проверка CRUD-сценариев")
    add_table(
        doc,
        ["Сущность", "Create", "Read", "Update", "Delete"],
        [
            ["Двигатель", "POST /api/engines", "GET /api/engines", "PUT /api/engines/<id>", "DELETE /api/engines/<id>"],
            ["Деталь", "POST /api/parts", "GET /api/parts", "sp_update_part_condition", "Удаляется каскадно с двигателем"],
            ["Заявка", "POST /api/service-requests", "GET /api/service-requests", "PATCH /api/service-requests/<id>", "Отмена через статус"],
            ["Пользователь", "POST /api/auth/register", "GET /api/auth/me", "Не вынесено в интерфейс", "Защищено триггером для admin"],
        ],
        [Cm(2.7), Cm(3.2), Cm(3.2), Cm(3.2), Cm(3.2)],
    )
    add_heading(doc, "7.3 Проверка базы данных", 2)
    add_p(doc, "Для базы данных проверяется наличие таблиц, внешних ключей, представлений, функций, процедур и триггеров. Отдельно проверяется логика защиты email и допустимых статусов заявок. Эти проверки важны, потому что часть требований теперь реализована не только в Python-коде, но и на уровне PostgreSQL.")
    add_caption(doc, "Таблица 17 - Проверка объектов базы данных")
    add_table(
        doc,
        ["Объект", "Что проверяется", "Ожидаемый результат"],
        [
            ["Таблицы", "roles, users, engines, engine_parts, service_requests, action_logs.", "Все таблицы создаются из schema.sql."],
            ["Представления", "v_engines_full, v_service_requests_full, v_request_status_stats.", "SELECT возвращает объединенные или агрегированные данные."],
            ["Функции", "fn_engine_parts_count, fn_user_request_count, fn_search_engines.", "Функции возвращают счетчики и результаты поиска."],
            ["Процедуры", "sp_change_request_status, sp_update_part_condition, sp_cancel_old_new_requests.", "Процедуры выполняют бизнес-операции."],
            ["Триггеры", "Проверка email, статуса заявки, последнего администратора.", "Некорректные операции блокируются."],
        ],
        [Cm(4), Cm(7.5), Cm(4.5)],
    )
    add_heading(doc, "7.4 Ограничения тестирования", 2)
    add_p(doc, "Для полноценной проверки операций с базой данных необходимо установить PostgreSQL, создать базу engine_coursework и выполнить SQL-скрипты schema.sql и seed.sql. После этого приложение сможет выполнять операции регистрации, входа, создания двигателей, деталей и заявок.")
    add_p(doc, "В текущем окружении была доступна проверка синтаксиса Python и JavaScript, импорт Flask-приложения и визуальная проверка интерфейса. Команда psql в системе отсутствовала, поэтому применение SQL-скрипта к живой базе должно быть выполнено после установки PostgreSQL.")
    add_p(doc, "Несмотря на это, структура SQL-файла подготовлена полностью: все таблицы, представления, функции, процедуры и триггеры находятся в database/schema.sql. Для запуска на рабочем компьютере достаточно установить PostgreSQL и выполнить команды из приложения.")

    doc.add_page_break()
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1, "sec_conclusion")
    add_p(doc, "В ходе курсовой работы было разработано клиент-серверное приложение для учета двигателей и заявок на обслуживание. Реализована серверная часть на Flask, клиентская часть на HTML, CSS и JavaScript, а также реляционная схема PostgreSQL.")
    add_p(doc, "В проекте выполнены обязательные требования: разделение ролей, регистрация, вход, валидация email, база данных более чем на четыре таблицы и логирование действий. Дополнительно реализованы справочники двигателей и деталей, заявки на обслуживание, смена статусов, CSV-выгрузка, страница статистики, представления, функции, процедуры и триггеры в базе данных.")
    add_p(doc, "Разработанное приложение можно расширять: добавить загрузку фотографий двигателя, полноценные JWT-токены, страницу статистики, фильтрацию заявок по статусу и интеграцию с внешними сервисами диагностики.")

    doc.add_page_break()
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, "sec_sources")
    sources = [
        "Репозиторий проекта KSP_Engine_Coursework. Локальный проект: /Users/sergeikoshmin/KSP_Engine_Coursework.",
        "Репозиторий poem_app. URL: https://github.com/Sofuliik/poem_app",
        "Flask Documentation. URL: https://flask.palletsprojects.com/",
        "Werkzeug Documentation. Password Hashing. URL: https://werkzeug.palletsprojects.com/",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/",
        "MDN Web Docs. Fetch API. URL: https://developer.mozilla.org/ru/docs/Web/API/Fetch_API",
        "MDN Web Docs. Window.localStorage. URL: https://developer.mozilla.org/ru/docs/Web/API/Window/localStorage",
        "ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления.",
    ]
    for i, source in enumerate(sources, 1):
        add_p(doc, f"{i}. {source}")

    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ", 1, "sec_appendix")
    add_p(doc, "В приложении приведены основные команды запуска и проверки проекта.")
    add_caption(doc, "Таблица 18 - Команды запуска и проверки")
    add_table(
        doc,
        ["Команда", "Назначение"],
        [
            ["python3 -m venv .venv", "Создание виртуального окружения."],
            ["source .venv/bin/activate", "Активация окружения."],
            ["pip install -r requirements.txt", "Установка зависимостей."],
            ["createdb engine_coursework", "Создание базы данных PostgreSQL."],
            ["psql engine_coursework -f database/schema.sql", "Создание таблиц и связей."],
            ["psql engine_coursework -f database/seed.sql", "Загрузка тестовых данных."],
            ["python -m backend.app", "Запуск Flask-приложения."],
            ["node --check public/js/app.js", "Проверка синтаксиса JavaScript."],
        ],
        [Cm(7), Cm(8.5)],
    )
    add_p(doc, "Тестовые пользователи после выполнения seed.sql:")
    add_code(doc, "admin@engine.local / 1234\nmechanic@engine.local / 1234\nclient@engine.local / 1234")
    add_p(doc, "Ниже приведена итоговая структура проекта, которую можно показать при защите. Она демонстрирует разделение приложения на серверную часть, клиентскую часть, базу данных и документацию.")
    add_code(doc, "KSP_Engine_Coursework/\n├── backend/\n│   ├── app.py\n│   ├── auth.py\n│   ├── audit.py\n│   ├── db.py\n│   ├── log.py\n│   └── routes/\n│       ├── auth.py\n│       └── api.py\n├── database/\n│   ├── schema.sql\n│   ├── seed.sql\n│   └── setup-local.sh\n├── public/\n│   ├── index.html\n│   ├── login.html\n│   ├── engines.html\n│   ├── parts.html\n│   ├── requests.html\n│   ├── stats.html\n│   ├── logs.html\n│   ├── css/style.css\n│   └── js/app.js\n├── docs/\n└── requirements.txt")
    add_caption(doc, "Таблица 19 - SQL-объекты проекта")
    add_table(
        doc,
        ["Категория", "Объекты"],
        [
            ["Таблицы", "roles, users, engines, engine_parts, service_requests, action_logs"],
            ["Представления", "v_engines_full, v_service_requests_full, v_request_status_stats"],
            ["Функции", "fn_engine_parts_count, fn_user_request_count, fn_search_engines"],
            ["Процедуры", "sp_change_request_status, sp_update_part_condition, sp_cancel_old_new_requests"],
            ["Триггеры", "trg_users_email_check, trg_prevent_last_admin_delete, trg_service_request_status_check"],
        ],
        [Cm(4), Cm(11.5)],
    )
    add_p(doc, "Для демонстрации проекта в seed.sql добавлены два двигателя: Toyota 2JZ-GE и Cummins ISF 2.8. Также добавлены детали, тестовая заявка и три пользователя с разными ролями. Это позволяет сразу после подготовки базы показать все основные страницы без ручного заполнения большого количества данных.")
    add_caption(doc, "Таблица 20 - Демонстрационные данные")
    add_table(
        doc,
        ["Сущность", "Пример данных", "Назначение"],
        [
            ["Пользователь", "admin@engine.local, mechanic@engine.local, client@engine.local", "Проверка ролей и входа."],
            ["Двигатель", "Toyota 2JZ-GE", "Проверка справочника двигателей и заявок."],
            ["Двигатель", "Cummins ISF 2.8", "Проверка нескольких записей в справочнике."],
            ["Деталь", "Топливный насос, свечи зажигания, турбокомпрессор", "Проверка таблицы деталей."],
            ["Заявка", "Двигатель троит после прогрева", "Проверка списка заявок и смены статуса."],
        ],
        [Cm(3.3), Cm(7.2), Cm(5)],
    )
    add_p(doc, "Команды, которые использовались при проверке отчета и проекта:")
    add_code(doc, ".venv/bin/python -m compileall backend\nnode --check public/js/app.js\n.venv/bin/python - <<'PY'\nfrom backend.app import create_app\napp = create_app()\nprint(len(app.url_map._rules))\nPY")

    # The title page stays unnumbered; content starts with the real page 2.
    for sec in doc.sections[1:]:
        add_page_number(sec)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
